import os
from docx import Document

class FileParsingError(Exception):
    """Raised when PDF or DOCX text extraction fails."""
    pass

class LLMProcessingError(Exception):
    """Raised when API operations fail."""
    pass

def init_workspace():
    """Ensure template directories and output folders exist."""
    os.makedirs("templates", exist_ok=True)
    os.makedirs("output", exist_ok=True)
    
    template_path = "templates/ats_template.docx"
    
    # Agar purani galat template file pehle se bani hai, toh use hata dete hain
    if os.path.exists(template_path):
        try:
            os.remove(template_path)
        except Exception:
            pass
            
    # Nayi correct syntax wali template file generate karte hain
    doc = Document()
    doc.add_heading("{{name}}", 0)
    doc.add_paragraph("Email: {{email}} | Phone: {{phone}}")
    
    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph("{{summary}}")
    
    doc.add_heading("Skills Matrix", level=1)
    doc.add_paragraph("{{skills_formatted}}")
    
    doc.add_heading("Professional Experience", level=1)
    # Sahi DocxTpl Loop Syntax: {% for ... %} aur {% endfor %}
    doc.add_paragraph("{% for item in experience %}")
    doc.add_heading("{{item.title}} - {{item.company}} ({{item.duration}})", level=2)
    doc.add_paragraph("{% for bullet in item.bullets %}")
    doc.add_paragraph("• {{bullet}}")
    doc.add_paragraph("{% endfor %}")
    doc.add_paragraph("{% endfor %}")
    
    doc.save(template_path)
